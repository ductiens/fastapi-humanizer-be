import fitz # PyMuPDF
import docx
from io import BytesIO

async def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        # PyMuPDF is much faster and more reliable than pdfplumber
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        for page in pdf:
            extracted = page.get_text()
            if extracted:
                text += extracted + "\n"
        pdf.close()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        raise Exception("Không thể đọc định dạng của file PDF này. File có thể bị hỏng hoặc mã hóa.")
    return text.strip()

async def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)

async def create_docx_from_text(text: str) -> BytesIO:
    doc = docx.Document()
    # Handle newlines for correct docx output
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
